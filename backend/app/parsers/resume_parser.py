"""
Resume parser module.
Extracts structured data from PDF and DOCX resume files:
- Candidate name (via fast structural/heuristic extraction + spaCy fallback)
- Contact info: email, phone
- Skills
- Work experience & internships (with date range & duration calculation)
- Education background (degree, institution, graduation year)
- Estimated total years of experience
"""
from __future__ import annotations

import logging
import os
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from dateutil import parser as dateutil_parser
from dateutil.relativedelta import relativedelta

logger = logging.getLogger(__name__)

# Lazy-loaded spaCy model container
_NLP_MODEL = None


def _get_nlp():
    """
    Lazy load spaCy model only when needed.
    """
    global _NLP_MODEL
    if _NLP_MODEL is None:
        try:
            import spacy  # noqa: PLC0415
            try:
                _NLP_MODEL = spacy.load("en_core_web_sm")
            except OSError:
                logger.info("Downloading en_core_web_sm...")
                from spacy.cli import download  # noqa: PLC0415
                download("en_core_web_sm")
                _NLP_MODEL = spacy.load("en_core_web_sm")
        except Exception as exc:
            logger.warning("spaCy model load failed, falling back to regex: %s", exc)
            _NLP_MODEL = None
    return _NLP_MODEL


# ─────────────────────────────────────────────────────────────────────────────
# Regex patterns
# ─────────────────────────────────────────────────────────────────────────────
_RE_EMAIL = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
_RE_PHONE = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?(?:\(?\d{3}\)?[\s\-.]?)?\d{3}[\s\-.]?\d{4}"
)
_RE_URL = re.compile(r"https?://\S+|www\.\S+")

# Section headers — ordered by priority (more specific first)
_SECTION_HEADERS = {
    "experience": re.compile(
        r"^\s*(internships?(\s+experience)?|work\s+experience|professional\s+experience|employment(\s+history)?|experience|work\s+history)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "education": re.compile(
        r"^\s*(education(\s+background)?|academic\s+(background|qualifications?|history)|qualifications?)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "skills": re.compile(
        r"^\s*(technical\s+skills|key\s+skills|skills?\s*(summary|set)?|competencies|expertise|technologies|tools)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "summary": re.compile(
        r"^\s*((brief\s+)?summary|profile|objective|about\s+me|professional\s+summary|executive\s+summary)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "certifications": re.compile(
        r"^\s*((assessments\s*/\s*)?certifications?|assessments?|certificates?|licenses?|courses?)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
    "projects": re.compile(
        r"^\s*(projects?|key\s+projects?|personal\s+projects?|academic\s+projects?|portfolio)\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    ),
}

# Date-range pattern: "26 Dec, 2024 - 31 Jan, 2025", "Jan 2020 – Present", "03/2018 - 06/2021", etc.
_RE_DATE_RANGE = re.compile(
    r"(?P<start>"
    r"(?:\d{1,2}\s+)?(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?"
    r"|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
    r"[\s,\.]+\d{4}"
    r"|\d{1,2}[/\-]\d{4}"
    r"|\d{4}"
    r")"
    r"\s*(?:–|—|-|to)\s*"
    r"(?P<end>"
    r"(?:\d{1,2}\s+)?(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?"
    r"|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
    r"[\s,\.]+\d{4}"
    r"|\d{1,2}[/\-]\d{4}"
    r"|\d{4}"
    r"|present|current|now"
    r")",
    re.IGNORECASE,
)

_RE_DEGREE = re.compile(
    r"\b(b\.?s\.?|b\.?e\.?|b\.?tech|b\.?sc|bachelor|m\.?s\.?|m\.?e\.?|m\.?tech|"
    r"m\.?sc|master|ph\.?d|doctorate|associate|diploma|high\s+school|10th|12th|secondary)\b",
    re.IGNORECASE,
)
_RE_YEAR = re.compile(r"\b(?:19|20)\d{2}\b")
_RE_INSTITUTION = re.compile(
    r"\b(university|college|institute|school|academy|polytechnic|campus|faculty)\b",
    re.IGNORECASE,
)

_RE_JOB_TITLES = re.compile(
    r"\b(intern|internship|developer|engineer|analyst|manager|consultant|specialist|lead|"
    r"architect|scientist|associate|assistant|designer|programmer|administrator)\b",
    re.IGNORECASE,
)

KNOWN_TECH_SKILLS = [
    "Python", "Java", "Core Java", "Java Swing", "AWT", "Socket Programming", "OOP", "Object Oriented Programming",
    "Data Structures", "Algorithms", "Data Structures & Algorithms", "C", "C++", "C#", "JavaScript", "TypeScript",
    "HTML", "HTML5", "CSS", "CSS3", "PHP", "WordPress", "Plugin Integration", "React", "Next.js", "Node.js",
    "SQL", "MySQL", "PostgreSQL", "MongoDB", "FastAPI", "Django", "Flask", "Spring Boot",
    "Data Science", "Data Analysis", "Machine Learning", "AI/ML", "Pandas", "Scikit-learn", "NumPy", "TensorFlow", "PyTorch",
    "REST API", "REST API Development", "API Testing", "Postman", "Postman Scripting", "Cloud Security", "Cloud Computing",
    "Oracle Cloud Infrastructure", "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Git", "GitHub", "Linux"
]


# ─────────────────────────────────────────────────────────────────────────────
# Data containers
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class ParsedResume:
    full_name: str = ""
    email: str = ""
    phone: str = ""
    skills: list[str] = field(default_factory=list)
    experience: list[dict[str, str]] = field(default_factory=list)
    education: list[dict[str, str]] = field(default_factory=list)
    total_experience_years: float = 0.0
    parse_warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "full_name": self.full_name,
            "email": self.email,
            "phone": self.phone,
            "skills": self.skills,
            "experience": self.experience,
            "education": self.education,
            "total_experience_years": self.total_experience_years,
            "parse_warnings": self.parse_warnings,
        }


# ─────────────────────────────────────────────────────────────────────────────
# File text extractors
# ─────────────────────────────────────────────────────────────────────────────

def _extract_text_pdf(file_path: Path) -> str:
    """
    Extract text from a PDF.
    Uses pypdfium2 for ultra-fast C++ extraction (<10ms),
    falling back to pdfplumber if necessary.
    """
    # 1. Fast path: pypdfium2
    try:
        import pypdfium2 as pdfium  # noqa: PLC0415

        pdf = pdfium.PdfDocument(file_path)
        try:
            pages_text: list[str] = []
            for page in pdf:
                textpage = page.get_textpage()
                text = textpage.get_text_range()
                if text and text.strip():
                    pages_text.append(text.strip())
            full_text = "\n\n".join(pages_text)
            if full_text.strip():
                return full_text
        finally:
            pdf.close()
    except Exception as exc:
        logger.debug("pypdfium2 extraction skipped/failed: %s, falling back to pdfplumber", exc)

    # 2. Fallback: pdfplumber fast stream extraction
    try:
        import pdfplumber  # noqa: PLC0415

        pages: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text(layout=False) or page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)
    except Exception as exc:
        logger.warning("pdfplumber extraction failed: %s", exc)
        return ""


def _extract_text_docx(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx  # noqa: PLC0415

        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs)
    except Exception as exc:
        logger.warning("DOCX extraction failed for %s: %s", file_path, exc)
        return ""


def extract_raw_text(file_path: str | Path) -> str:
    """
    Extract raw text from a PDF or DOCX resume file.
    Raises ValueError for unsupported formats.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _extract_text_docx(path)
    else:
        raise ValueError(f"Unsupported file format '{suffix}'. Supported: .pdf, .docx")


# ─────────────────────────────────────────────────────────────────────────────
# Section segmentation
# ─────────────────────────────────────────────────────────────────────────────

def _segment_sections(text: str) -> dict[str, str]:
    """
    Split the resume text into named sections using header regexes.
    Returns a dict mapping section name -> section text.
    The text before the first recognised header is stored under 'header'.
    """
    matches: list[tuple[int, int, str]] = []
    for sec_name, pattern in _SECTION_HEADERS.items():
        for m in pattern.finditer(text):
            matches.append((m.start(), m.end(), sec_name))

    matches.sort(key=lambda x: x[0])

    if not matches:
        return {"header": text}

    sections: dict[str, str] = {}
    sections["header"] = text[: matches[0][0]].strip()

    for i, (start, end, sec_name) in enumerate(matches):
        next_start = matches[i + 1][0] if i + 1 < len(matches) else len(text)
        content = text[end:next_start].strip()
        if sec_name in sections:
            sections[sec_name] += "\n\n" + content
        else:
            sections[sec_name] = content

    return sections


# ─────────────────────────────────────────────────────────────────────────────
# Contact extraction
# ─────────────────────────────────────────────────────────────────────────────

def _extract_email(text: str) -> str:
    match = _RE_EMAIL.search(text)
    return match.group(0).strip() if match else ""


def _extract_phone(text: str) -> str:
    match = _RE_PHONE.search(text)
    return match.group(0).strip() if match else ""


# ─────────────────────────────────────────────────────────────────────────────
# Name extraction (Fast Heuristics + spaCy Fallback)
# ─────────────────────────────────────────────────────────────────────────────

def _extract_name(header_text: str, warnings: list[str]) -> str:
    """
    Extract candidate name.
    1. Fast heuristic: First non-empty, non-contact line of header with 2-4 capitalized words.
    2. Fallback: spaCy PERSON entity recognition.
    """
    lines = [l.strip() for l in header_text.splitlines() if l.strip()]
    if not lines:
        return ""

    for line in lines[:5]:
        if _RE_EMAIL.search(line) or _RE_PHONE.search(line) or _RE_URL.search(line):
            continue
        cleaned = re.sub(r"[^\w\s\.\-]", "", line).strip()
        words = cleaned.split()
        if 2 <= len(words) <= 4:
            if all(w[0].isupper() or w.isupper() for w in words if w):
                return cleaned

    nlp = _get_nlp()
    if nlp is not None:
        try:
            doc = nlp(header_text[:400])
            for ent in doc.ents:
                if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                    return ent.text.strip()
        except Exception as exc:
            logger.debug("spaCy PERSON extraction failed: %s", exc)

    if lines:
        first_line = re.sub(r"[^\w\s\.\-]", "", lines[0]).strip()
        if first_line and len(first_line.split()) <= 5:
            return first_line

    warnings.append("Could not reliably extract candidate name.")
    return ""


# ─────────────────────────────────────────────────────────────────────────────
# Skills extraction
# ─────────────────────────────────────────────────────────────────────────────

def _extract_skills(skills_text: str, full_text: str = "") -> list[str]:
    """
    Extract skills from:
    1. Dedicated Skills section
    2. 'Key Skills:' lines across projects, internships, and certifications
    3. Technical keyword dictionary matching
    """
    extracted: list[str] = []
    seen: set[str] = set()

    def add_skill(s: str):
        clean = s.strip("()[]{}'\".,;:- ")
        k = clean.lower()
        if clean and len(clean) >= 2 and k not in seen:
            seen.add(k)
            extracted.append(clean)

    # 1. From skills section
    if skills_text:
        for token in re.split(r"[•·▪▸►\-\*\|,;\n/]", skills_text):
            tok = token.strip()
            if ":" in tok:
                tok = tok.split(":", 1)[1].strip()
            if tok and len(tok) <= 40:
                add_skill(tok)

    # 2. From 'Key Skills:' lines
    if full_text:
        for line in full_text.splitlines():
            if re.search(r"^\s*key\s+skills\s*:", line, re.IGNORECASE):
                parts = re.split(r"^\s*key\s+skills\s*:\s*", line, flags=re.IGNORECASE)[1]
                tokens = [t.strip() for t in re.split(r"[,•|;]|\s{2,}", parts) if t.strip()]
                for t in tokens:
                    add_skill(t)

        # 3. Known tech keyword lookup
        full_lower = f" {full_text.lower()} "
        for sk in KNOWN_TECH_SKILLS:
            pattern = r"\b" + re.escape(sk.lower()) + r"\b"
            if re.search(pattern, full_lower):
                add_skill(sk)

    return extracted


# ─────────────────────────────────────────────────────────────────────────────
# Experience extraction
# ─────────────────────────────────────────────────────────────────────────────

def _parse_date(date_str: str) -> Any:
    """Parse a fuzzy date string into a datetime; return None on failure."""
    try:
        return dateutil_parser.parse(date_str, fuzzy=True, default=dateutil_parser.parse("1900-01-01"))
    except Exception:
        return None


def _compute_duration_years(start_str: str, end_str: str) -> float:
    """Return fractional years between two date strings. Returns 0 on failure."""
    from datetime import datetime, timezone  # noqa: PLC0415

    start = _parse_date(start_str)
    end_lower = end_str.strip().lower()
    if end_lower in ("present", "current", "now"):
        end = datetime.now(timezone.utc).replace(tzinfo=None)
    else:
        end = _parse_date(end_str)

    if start is None or end is None or end < start:
        return 0.0

    delta = relativedelta(end, start)
    return round(delta.years + delta.months / 12, 2)


def _extract_experience(exp_text: str, warnings: list[str]) -> tuple[list[dict[str, str]], float]:
    """
    Parse the experience section into a list of job/internship entries with fast structural extraction.
    Each entry: {title, company, duration, description}
    """
    if not exp_text:
        return [], 0.0

    entries: list[dict[str, str]] = []
    total_years = 0.0

    # Split into blocks on empty lines
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in exp_text.splitlines():
        stripped = line.strip()
        if not stripped:
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(stripped)
    if current:
        blocks.append(current)

    for block in blocks:
        if not block:
            continue

        block_text = "\n".join(block)
        entry: dict[str, str] = {
            "title": "",
            "company": "",
            "duration": "",
            "description": "",
        }

        # ── Duration ────────────────────────────────────────────────────────
        date_match = _RE_DATE_RANGE.search(block_text)
        if date_match:
            start_str = date_match.group("start")
            end_str = date_match.group("end")
            entry["duration"] = date_match.group(0).strip()
            total_years += _compute_duration_years(start_str, end_str)
            block_lines = [
                l.replace(date_match.group(0), "").strip()
                for l in block
                if l.replace(date_match.group(0), "").strip()
            ]
        else:
            block_lines = [l.strip() for l in block if l.strip()]

        if not block_lines:
            continue

        line1 = block_lines[0]
        line2 = block_lines[1] if len(block_lines) > 1 else ""

        # Title & Company Detection
        if line2 and _RE_JOB_TITLES.search(line2) and not _RE_JOB_TITLES.search(line1):
            entry["title"] = line2
            entry["company"] = line1.split("|")[0].strip()
            desc_start = 2
        elif "|" in line1:
            parts = [p.strip() for p in line1.split("|") if p.strip()]
            entry["company"] = parts[0]
            entry["title"] = line2 if line2 else (parts[1] if len(parts) > 1 else parts[0])
            desc_start = 2 if line2 else 1
        elif " at " in line1.lower():
            parts = re.split(r"\s+at\s+", line1, flags=re.IGNORECASE, maxsplit=1)
            entry["title"] = parts[0].strip()
            entry["company"] = parts[1].strip()
            desc_start = 1
        else:
            entry["title"] = line1
            entry["company"] = line2
            desc_start = 2 if line2 else 1

        entry["description"] = " ".join(block_lines[desc_start:])
        if entry["title"] or entry["company"]:
            entries.append(entry)
        else:
            warnings.append(f"Could not parse an experience block: {block_text[:60]!r}")

    return entries, round(total_years, 2)


# ─────────────────────────────────────────────────────────────────────────────
# Education extraction
# ─────────────────────────────────────────────────────────────────────────────

def _extract_education(edu_text: str, warnings: list[str]) -> list[dict[str, str]]:
    """
    Parse the education section into a list of education entries.
    Each entry: {degree, institution, year}
    """
    if not edu_text:
        return []

    entries: list[dict[str, str]] = []

    blocks: list[list[str]] = []
    current: list[str] = []
    for line in edu_text.splitlines():
        stripped = line.strip()
        if not stripped:
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(stripped)
    if current:
        blocks.append(current)

    for block in blocks:
        if not block:
            continue
        entry: dict[str, str] = {"degree": "", "institution": "", "year": ""}
        block_text = "\n".join(block)

        # Year
        years = _RE_YEAR.findall(block_text)
        if years:
            entry["year"] = years[-1]

        # Degree & Institution
        for line in block:
            if _RE_DEGREE.search(line) and not entry["degree"]:
                entry["degree"] = line.strip()
            if _RE_INSTITUTION.search(line) and not entry["institution"]:
                entry["institution"] = _RE_YEAR.sub("", line).strip(" -–|,;").strip()

        if not entry["institution"] and block:
            entry["institution"] = _RE_YEAR.sub("", block[0]).strip(" -–|,;").strip()
            if not entry["degree"] and len(block) > 1:
                entry["degree"] = block[1].strip()

        if entry["degree"] or entry["institution"]:
            entries.append(entry)
        else:
            warnings.append(f"Could not parse an education block: {block_text[:60]!r}")

    return entries


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def parse_resume_with_raw_text(file_path: str | Path) -> tuple[dict[str, Any], str]:
    """
    Parse a PDF or DOCX resume file and return both structured dict and extracted raw text.
    Eliminates redundant text extraction passes.
    """
    file_path = Path(file_path)
    result = ParsedResume()

    # ── 1. Raw text extraction (Single pass) ─────────────────────────────────
    try:
        raw_text = extract_raw_text(file_path)
    except Exception as exc:
        result.parse_warnings.append(f"Text extraction failed: {exc}")
        return result.to_dict(), ""

    if not raw_text.strip():
        result.parse_warnings.append("Extracted text is empty — the file may be image-based or corrupt.")
        return result.to_dict(), ""

    # ── 2. Section segmentation ─────────────────────────────────────────────
    try:
        sections = _segment_sections(raw_text)
    except Exception as exc:
        result.parse_warnings.append(f"Section segmentation failed: {exc}")
        sections = {"header": raw_text}

    header_text = sections.get("header", raw_text[:500])

    # ── 3. Contact info ─────────────────────────────────────────────────────
    try:
        result.email = _extract_email(raw_text)
    except Exception as exc:
        result.parse_warnings.append(f"Email extraction error: {exc}")

    try:
        result.phone = _extract_phone(raw_text)
    except Exception as exc:
        result.parse_warnings.append(f"Phone extraction error: {exc}")

    # ── 4. Name ─────────────────────────────────────────────────────────────
    try:
        result.full_name = _extract_name(header_text, result.parse_warnings)
    except Exception as exc:
        result.parse_warnings.append(f"Name extraction error: {exc}")

    # ── 5. Skills ───────────────────────────────────────────────────────────
    try:
        result.skills = _extract_skills(sections.get("skills", ""), raw_text)
    except Exception as exc:
        result.parse_warnings.append(f"Skills extraction error: {exc}")

    # ── 6. Experience ───────────────────────────────────────────────────────
    try:
        result.experience, result.total_experience_years = _extract_experience(
            sections.get("experience", ""),
            result.parse_warnings,
        )
    except Exception as exc:
        result.parse_warnings.append(f"Experience extraction error: {exc}")

    # ── 7. Education ────────────────────────────────────────────────────────
    try:
        result.education = _extract_education(
            sections.get("education", ""),
            result.parse_warnings,
        )
    except Exception as exc:
        result.parse_warnings.append(f"Education extraction error: {exc}")

    return result.to_dict(), raw_text


def parse_resume(file_path: str | Path) -> dict[str, Any]:
    """
    Parse a PDF or DOCX resume file and return a structured dict.
    """
    parsed_dict, _ = parse_resume_with_raw_text(file_path)
    return parsed_dict
